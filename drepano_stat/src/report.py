"""Génération du rapport analytique Word de DrepanoStat."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Mapping, Sequence

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib.figure import Figure

from src.display_labels import apply_display_labels, display_label
from src.plots import save_figure_to_bytes


def _format_value(value: Any) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.4g}"
    return str(display_label(value))


def _add_dataframe(document: Document, title: str, dataframe: pd.DataFrame) -> None:
    document.add_heading(title, level=2)
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        document.add_paragraph("Aucun résultat disponible pour cette section.")
        return
    displayed = apply_display_labels(dataframe)
    table = document.add_table(rows=1, cols=len(displayed.columns))
    table.style = "Table Grid"
    table.autofit = True
    for cell, column in zip(table.rows[0].cells, displayed.columns, strict=True):
        cell.text = str(column)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in displayed.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = _format_value(value)
    document.add_paragraph()


def generate_word_report(
    study_info: Mapping[str, Any],
    validation_messages: Sequence[str],
    descriptive_table: pd.DataFrame,
    ranking_table: pd.DataFrame,
    group_ranking_table: pd.DataFrame,
    glm_results_table: pd.DataFrame,
    pairwise_results_table: pd.DataFrame,
    figures: Mapping[str, Figure],
) -> bytes:
    """Génère en mémoire un rapport Word analytique et prudent."""
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    title = document.add_heading("DrepanoStat — Rapport analytique", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "Synthèse descriptive et statistique des données expérimentales"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("1. Informations de l’étude", level=1)
    for key, value in study_info.items():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{key} : ").bold = True
        if isinstance(value, (list, tuple, set)):
            value = ", ".join(str(display_label(item)) for item in value)
        paragraph.add_run(_format_value(value))

    document.add_heading("2. Contrôle des données", level=1)
    if validation_messages:
        document.add_paragraph(
            "Le fichier a permis la poursuite de l’analyse. Les messages suivants ont été relevés :"
        )
        for message in validation_messages:
            document.add_paragraph(str(display_label(message)), style="List Bullet")
    else:
        document.add_paragraph(
            "Aucune erreur bloquante ni aucun avertissement n’a été relevé lors de la validation."
        )

    document.add_heading("3. Méthodologie statistique", level=1)
    document.add_paragraph(
        "Les proportions ont été calculées séparément pour chaque répétition à partir des "
        "comptages N (globules rouges normaux) et D (globules rouges drépanocytaires). "
        "Les comparaisons ont été ajustées par modèles linéaires généralisés binomiaux "
        "avec lien logit, en utilisant directement la réponse groupée [N, D]."
    )
    document.add_paragraph(
        "Les conditions d’extraits ont été comparées au Témoin véhicule. Les comparaisons "
        "entre groupes ont été réalisées à dilution identique. Les p-values des familles "
        "de comparaisons ont été corrigées par la méthode de Holm. Les odds ratios sont "
        "présentés avec leurs intervalles de confiance à 95 %."
    )

    document.add_heading("4. Résultats descriptifs", level=1)
    _add_dataframe(document, "Tableau descriptif par condition", descriptive_table)
    _add_dataframe(
        document,
        "Classement des conditions par proportion observée de drépanocytes",
        ranking_table,
    )
    _add_dataframe(document, "Classement descriptif global des groupes", group_ranking_table)

    document.add_heading("5. GLM vs Témoin véhicule", level=1)
    _add_dataframe(document, "Résultats des odds ratios", glm_results_table)

    document.add_heading("6. Comparaisons entre groupes", level=1)
    _add_dataframe(document, "Comparaisons à dilution identique", pairwise_results_table)

    document.add_heading("7. Figures", level=1)
    if figures:
        for figure_name, figure in figures.items():
            document.add_heading(str(figure_name), level=2)
            image = BytesIO(save_figure_to_bytes(figure, format="png", dpi=300))
            document.add_picture(image, width=Inches(6.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph("Aucune figure disponible.")

    document.add_heading("8. Note méthodologique", level=1)
    document.add_paragraph(
        "Ce rapport décrit les effets observés dans les données et les résultats des modèles "
        "statistiques dans les conditions expérimentales étudiées. Le classement indique la "
        "condition présentant la plus faible proportion observée de drépanocytes ; il ne "
        "constitue pas à lui seul une conclusion biologique. Les résultats dépendent notamment "
        "du nombre de répétitions, du nombre de cellules comptées et de la variabilité "
        "expérimentale. Leur interprétation finale est laissée à l’utilisateur et doit tenir "
        "compte du protocole expérimental ainsi que d’éventuels essais complémentaires."
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()

